"""Text extraction from uploaded dossier files.

PDF text layer via pypdf, Word via python-docx, images via Tesseract OCR
(Vietnamese + English traineddata when installed). Plain text is decoded
directly.

Every path reports honestly: a scanned PDF with no text layer, a missing
Tesseract binary, or an unsupported type produce an explicit error rather
than an empty string that would read as "extracted nothing". A legal tool
that silently returns the wrong dossier text is worse than one that refuses.
"""

import io
import os
import shutil
from typing import Any, Dict

MAX_BYTES = 25 * 1024 * 1024  # 25 MB per file

SUPPORTED = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "text",
    ".md": "text",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".webp": "image",
    ".tif": "image",
    ".tiff": "image",
}


def _ok(text: str, method: str, **extra: Any) -> Dict[str, Any]:
    return {"ok": True, "text": text, "method": method, "chars": len(text), **extra}


def _err(code: str, message: str) -> Dict[str, Any]:
    return {"ok": False, "text": "", "error_code": code, "error": message}


def kind_for(filename: str) -> str:
    return SUPPORTED.get(os.path.splitext(filename or "")[1].lower(), "")


def tesseract_available() -> bool:
    if shutil.which("tesseract"):
        return True
    return bool(os.getenv("TESSERACT_CMD") and os.path.exists(os.getenv("TESSERACT_CMD", "")))


def extract(
    filename: str,
    data: bytes,
    classification: str = None,
    allow_external_ocr: bool = False,
) -> Dict[str, Any]:
    """Extract text from one uploaded file. Never raises."""
    if not data:
        return _err("empty_file", "Tệp rỗng.")
    if len(data) > MAX_BYTES:
        return _err("too_large", f"Tệp vượt quá {MAX_BYTES // (1024 * 1024)} MB.")

    kind = kind_for(filename)
    if not kind:
        ext = os.path.splitext(filename or "")[1] or "?"
        return _err("unsupported", f"Định dạng {ext} chưa được hỗ trợ.")

    try:
        if kind == "pdf":
            return _extract_pdf(data)
        if kind == "docx":
            return _extract_docx(data)
        if kind == "image":
            return _extract_image(data, classification, allow_external_ocr)
        return _extract_text(data)
    except Exception as err:
        return _err("extract_failed", f"Không đọc được tệp: {err}")


def _extract_text(data: bytes) -> Dict[str, Any]:
    for encoding in ("utf-8", "utf-16", "cp1258", "latin-1"):
        try:
            return _ok(data.decode(encoding).strip(), "plain-text")
        except UnicodeDecodeError:
            continue
    return _err("decode_failed", "Không giải mã được nội dung văn bản.")


def _extract_pdf(data: bytes) -> Dict[str, Any]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception:
            return _err("encrypted", "PDF được bảo vệ bằng mật khẩu.")

    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")

    text = "\n".join(pages).strip()
    if text:
        return _ok(text, "pdf-text-layer", pages=len(reader.pages))

    # No text layer — a scanned document. Say so instead of returning "".
    return _err(
        "pdf_no_text_layer",
        f"PDF này là bản scan ({len(reader.pages)} trang), không có lớp văn bản. "
        "Hãy tải lên ảnh từng trang để OCR, hoặc dùng bản PDF có text.",
    )


def _extract_docx(data: bytes) -> Dict[str, Any]:
    import docx

    document = docx.Document(io.BytesIO(data))
    blocks = [p.text for p in document.paragraphs if p.text.strip()]

    # Legal annexes carry a lot of their content in tables.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    text = "\n".join(blocks).strip()
    if not text:
        return _err("docx_empty", "Tệp Word không có nội dung văn bản.")
    return _ok(text, "docx")


def _extract_image(
    data: bytes,
    classification: str = None,
    allow_external_ocr: bool = False,
) -> Dict[str, Any]:
    from app.domains.legal_assistant.cloud_ocr import OCRProviderError, configured_chain

    # The chain tries local Tesseract first. Cloud providers are appended only
    # after explicit consent and are still blocked for secret classifications.
    chain = configured_chain(allow_external_ocr=allow_external_ocr)
    try:
        result = chain.recognize(data, classification=classification)
        return _ok(
            result.text,
            result.method,
            ocr_provider=result.provider,
            fallback_errors=chain.last_errors,
        )
    except OCRProviderError as err:
        if not tesseract_available() and not allow_external_ocr:
            return _err(
                "ocr_unavailable",
                "Chưa có OCR local và chưa bật Cloud OCR. Cài Tesseract hoặc bật Cloud OCR có xác nhận.",
            )
        return _err("ocr_failed", str(err))
